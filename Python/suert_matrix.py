#!/usr/bin/python
# ^_^ coding:utf8 ^_^
# sudo pip install python-docx
import docx
import random
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT


def gen_suert_matrix(length):
    numbers = [i for i in range(1, length**2+1)]
    suert_matrix_linearity = random.sample(numbers, length**2)
    suert_matrix = list()
    for i in range(length):
        suert_matrix.append(suert_matrix_linearity[i*length: (i+1)*length])
    return suert_matrix
    
def add_table(doc, matrix):    
    table = doc.add_table(rows=len(matrix), cols=len(matrix), style="Light List Accent 5")
    table.style.font.size = Pt(18)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in range(len(matrix)):
        #table.rows[row].height = Pt(22)
        table.columns[row].width = Pt(30)
        for col in range(len(matrix[row])):
            table.cell(row, col).text = str(matrix[row][col])
            table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #table.cell(row, col).width = Pt(22)
            #table.cell(row, col).height = Pt(22)


def write_doc(matrixs, doc_path):
    doc = docx.Document()
    for matrix in matrixs:
        add_table(doc, matrix)
        doc.add_page_break()
    doc.save(doc_path)


if __name__ == '__main__':
    length = 5    # 方格边长
    count  = 1000 # 生成方格数量
    doc_path = './suert_matrix.doc'    # 输出文件路径


    matrixs = list()
    for i in range(count):
        matrixs.append(gen_suert_matrix(length))

    write_doc(matrixs, doc_path)

